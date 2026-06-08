#!/usr/bin/env python3
"""Convert a Markdown file to a .docx document.

Self-contained Markdown -> Word converter built on python-docx (no pandoc
binary required). Handles the subset of Markdown used by the project memo:
ATX headings, pipe tables, ordered/unordered (nested) lists, blockquotes,
fenced code blocks, horizontal rules, and inline **bold** / `code`.

Usage:
    python md_to_docx.py [input.md] [output.docx]

With no arguments it converts the presentation memo sitting next to this
script and writes a .docx with the same base name.
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit(
        "python-docx is not installed. Install it with:\n"
        "    python -m pip install python-docx"
    )

# Microsoft JhengHei ships with Windows zh-TW and renders Traditional Chinese.
CJK_FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
LIST_RE = re.compile(r"^(\s*)([-*+]|\d+\.)\s+(.*)$")
HR_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")


def set_style_east_asia(style, font=CJK_FONT):
    """Point a paragraph style's East-Asian font at a CJK-capable typeface.

    Latin glyphs keep the template default (Calibri); Word picks the
    East-Asian font per character, so mixed CN/EN text renders correctly.
    """
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_run_east_asia(run, font):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def tokenize(text, regex):
    """Split text into (chunk, matched?) pieces around regex matches (group 1)."""
    out = []
    last = 0
    for m in regex.finditer(text):
        if m.start() > last:
            out.append((text[last:m.start()], False))
        out.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        out.append((text[last:], False))
    return out


def add_inline(paragraph, text, force_bold=False):
    """Append text to a paragraph honoring **bold** and `code` markers.

    Bold is parsed first, then code within each bold/plain segment, so a
    bold span that contains an inline code fragment renders both styles.
    """
    for seg, is_bold in tokenize(text, BOLD_RE):
        for piece, is_code in tokenize(seg, CODE_RE):
            if piece == "":
                continue
            run = paragraph.add_run(piece)
            if is_bold or force_bold:
                run.bold = True
            if is_code:
                run.font.name = CODE_FONT
                set_run_east_asia(run, CODE_FONT)


def strip_inline(text):
    """Drop inline markers for contexts (headings) where styling is skipped."""
    return text.replace("**", "").replace("`", "")


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    pPr.append(pbdr)


def is_table_separator(line):
    line = line.strip()
    if "|" not in line or "-" not in line:
        return False
    cells = [c.strip() for c in line.strip("|").split("|")]
    cells = [c for c in cells if c != ""]
    return bool(cells) and all(re.fullmatch(r":?-+:?", c) for c in cells)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def list_style(doc, ordered, level):
    base = "List Number" if ordered else "List Bullet"
    name = base if level == 0 else f"{base} {level + 1}"
    try:
        doc.styles[name]
        return name
    except KeyError:
        return base


def convert(md_path, docx_path):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = Document()

    # Make the body + heading styles render Traditional Chinese.
    set_style_east_asia(doc.styles["Normal"])
    doc.styles["Normal"].font.size = Pt(11)
    for sname in ("Heading 1", "Heading 2", "Heading 3", "Heading 4",
                  "Title", "List Bullet", "List Number", "Quote"):
        try:
            set_style_east_asia(doc.styles[sname])
        except KeyError:
            pass

    n = len(lines)
    i = 0
    quote_buf = []

    def flush_quote():
        if not quote_buf:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_inline(p, "\n".join(quote_buf).strip())
        for run in p.runs:
            run.italic = True
        quote_buf.clear()

    while i < n:
        line = lines[i].rstrip()

        # Blockquote: accumulate consecutive '>' lines.
        if line.lstrip().startswith(">"):
            quote_buf.append(re.sub(r"^\s*>\s?", "", line))
            i += 1
            continue
        flush_quote()

        # Fenced code block.
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume the closing fence
            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl if cl else " ")
                run.font.name = CODE_FONT
                set_run_east_asia(run, CODE_FONT)
                run.font.size = Pt(10)
            continue

        # Blank line.
        if line.strip() == "":
            i += 1
            continue

        # Horizontal rule.
        if HR_RE.match(line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Pipe table (header row followed by a separator row).
        if line.strip().startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            header = split_row(line)
            ncols = len(header)
            body = []
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                body.append(split_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1 + len(body), cols=ncols)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            for c, htext in enumerate(header):
                cell = table.cell(0, c)
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], htext, force_bold=True)
            for r, row in enumerate(body, start=1):
                for c in range(ncols):
                    cell = table.cell(r, c)
                    cell.paragraphs[0].text = ""
                    add_inline(cell.paragraphs[0], row[c] if c < len(row) else "")
            doc.add_paragraph()  # spacer after the table
            i = j
            continue

        # Heading.
        m = HEADING_RE.match(line)
        if m:
            doc.add_heading(strip_inline(m.group(2)), level=min(len(m.group(1)), 4))
            i += 1
            continue

        # List item (nesting from leading spaces; 2 spaces = 1 level).
        m = LIST_RE.match(line)
        if m:
            level = min(len(m.group(1).replace("\t", "  ")) // 2, 2)
            ordered = bool(re.match(r"\d+\.", m.group(2)))
            p = doc.add_paragraph(style=list_style(doc, ordered, level))
            add_inline(p, m.group(3))
            i += 1
            continue

        # Plain paragraph.
        add_inline(doc.add_paragraph(), line)
        i += 1

    flush_quote()
    doc.save(str(docx_path))
    return docx_path


def main():
    here = Path(__file__).resolve().parent
    default_in = here / "presentation-notes" / "簡報備忘錄-技術問答.md"
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else default_in
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    if not md_path.exists():
        sys.exit(f"Input not found: {md_path}")
    out = convert(md_path, docx_path)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
