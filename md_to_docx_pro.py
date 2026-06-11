#!/usr/bin/env python3
"""Convert a Markdown report to a professional two-column .docx document.

Built on top of the project's md_to_docx.py converter, but tuned for a
formal written report:

  * A4 paper with tightened margins.
  * The leading title block (H1 + member metadata) stays full width; the
    body flows in two columns. The first '---' in the document marks that
    single -> two-column break.
  * Horizontal rules are never drawn -- every '---' is dropped (the first
    one only triggers the column break).
  * Justified body text, restyled headings/title, and the member table
    (originally trapped inside a '>' blockquote) rendered as a real table.

Usage:
    python md_to_docx_pro.py [input.md] [output.docx]
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.exit(
        "python-docx is not installed. Install it with:\n"
        "    python -m pip install python-docx"
    )

CJK_FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x38, 0x64)   # dark navy for headings
RULE_GRAY = RGBColor(0x66, 0x66, 0x66)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
LIST_RE = re.compile(r"^(\s*)([-*+]|\d+\.)\s+(.*)$")
HR_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")


def set_style_east_asia(style, font=CJK_FONT):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_run_east_asia(run, font):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_columns(section, num, space_twips=420):
    """Set a section's column count (and inter-column gap, in twips)."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:sep"), "0")  # no vertical separator line


def tokenize(text, regex):
    out = []
    last = 0
    for m in regex.finditer(text):
        if m.start() > last:
            out.append((text[last:m.start()], False))
        out.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        out.append((text[last:], False))
    return out


def add_inline(paragraph, text, force_bold=False):
    for seg, is_bold in tokenize(text, BOLD_RE):
        for piece, is_code in tokenize(seg, CODE_RE):
            if piece == "":
                continue
            run = paragraph.add_run(piece)
            if is_bold or force_bold:
                run.bold = True
            if is_code:
                run.font.name = CODE_FONT
                set_run_east_asia(run, CODE_FONT)


def strip_inline(text):
    return text.replace("**", "").replace("`", "")


def is_table_separator(line):
    line = line.strip()
    if "|" not in line or "-" not in line:
        return False
    cells = [c.strip() for c in line.strip("|").split("|")]
    cells = [c for c in cells if c != ""]
    return bool(cells) and all(re.fullmatch(r":?-+:?", c) for c in cells)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def list_style(doc, ordered, level):
    base = "List Number" if ordered else "List Bullet"
    name = base if level == 0 else f"{base} {level + 1}"
    try:
        doc.styles[name]
        return name
    except KeyError:
        return base


def render_table(doc, lines, i):
    """Render a pipe table starting at lines[i]; return the next index."""
    n = len(lines)
    header = split_row(lines[i])
    ncols = len(header)
    body = []
    j = i + 2
    while j < n and lines[j].strip().startswith("|"):
        body.append(split_row(lines[j]))
        j += 1
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    for c, htext in enumerate(header):
        cell = table.cell(0, c)
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], htext, force_bold=True)
    for r, row in enumerate(body, start=1):
        for c in range(ncols):
            cell = table.cell(r, c)
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], row[c] if c < len(row) else "")
    doc.add_paragraph()
    return j


def render_quote(doc, buf):
    """Render an accumulated blockquote, promoting any pipe table to a real table."""
    n = len(buf)
    i = 0
    while i < n:
        line = buf[i]
        if line.strip().startswith("|") and i + 1 < n and is_table_separator(buf[i + 1]):
            i = render_table(doc, buf, i)
            continue
        if line.strip() == "":
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, line.strip())
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RULE_GRAY
        i += 1


def style_document(doc):
    """Page setup + base/heading/title typography for a formal report."""
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(18)
    sec.right_margin = Mm(18)

    normal = doc.styles["Normal"]
    set_style_east_asia(normal)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for sname in ("Heading 1", "Heading 2", "Heading 3", "Heading 4",
                  "Title", "List Bullet", "List Number", "Quote"):
        try:
            set_style_east_asia(doc.styles[sname])
        except KeyError:
            pass

    heading_specs = {
        "Heading 1": (15, True),
        "Heading 2": (13, True),
        "Heading 3": (11.5, True),
        "Heading 4": (10.5, True),
    }
    for sname, (size, bold) in heading_specs.items():
        try:
            st = doc.styles[sname]
        except KeyError:
            continue
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.keep_with_next = True

    try:
        title = doc.styles["Title"]
        title.font.size = Pt(22)
        title.font.bold = True
        title.font.color.rgb = ACCENT
    except KeyError:
        pass


def convert(md_path, docx_path):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = Document()
    style_document(doc)
    set_columns(doc.sections[0], 1)   # title block stays full width

    n = len(lines)
    i = 0
    quote_buf = []
    column_break_done = False
    first_h1_done = False

    def flush_quote():
        if not quote_buf:
            return
        render_quote(doc, quote_buf)
        quote_buf.clear()

    while i < n:
        line = lines[i].rstrip()

        if line.lstrip().startswith(">"):
            quote_buf.append(re.sub(r"^\s*>\s?", "", line))
            i += 1
            continue
        flush_quote()

        # Fenced code block.
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            for cl in code_lines:
                p = doc.add_paragraph()
                # Diagram height is dominated by per-line paragraph
                # spacing, so code lines stack with none.
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(cl if cl else " ")
                run.font.name = CODE_FONT
                set_run_east_asia(run, CODE_FONT)
                run.font.size = Pt(9.5)
            if code_lines:
                # Restore a gap between the block and the next paragraph.
                p.paragraph_format.space_after = Pt(6)
            continue

        if line.strip() == "":
            i += 1
            continue

        # Horizontal rule: never drawn. The first one switches to two columns.
        if HR_RE.match(line):
            if not column_break_done:
                new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
                set_columns(new_sec, 2)
                column_break_done = True
            i += 1
            continue

        # Pipe table.
        if line.strip().startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            i = render_table(doc, lines, i)
            continue

        # Heading.
        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            text = strip_inline(m.group(2))
            if level == 1 and not first_h1_done:
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_h1_done = True
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        # List item.
        m = LIST_RE.match(line)
        if m:
            level = min(len(m.group(1).replace("\t", "  ")) // 2, 2)
            ordered = bool(re.match(r"\d+\.", m.group(2)))
            p = doc.add_paragraph(style=list_style(doc, ordered, level))
            add_inline(p, m.group(3))
            i += 1
            continue

        # Plain paragraph (justified).
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, line)
        i += 1

    flush_quote()
    doc.save(str(docx_path))
    return docx_path


def main():
    here = Path(__file__).resolve().parent
    default_in = here / "docs" / "期末書面報告.md"
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else default_in
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    if not md_path.exists():
        sys.exit(f"Input not found: {md_path}")
    out = convert(md_path, docx_path)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
